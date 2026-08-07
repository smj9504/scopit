from __future__ import annotations

"""
Scopit - PDF Editor Service

Handles PDF document CRUD, page operations (merge, reorder, delete, rotate),
annotation persistence, flattening, multi-image-to-PDF conversion, and the
company document library.

Storage: Abstracts file persistence via StorageBackend (local or R2).
Database: SQLAlchemy 2.0 SYNC (Session, not AsyncSession).
"""

import base64
import os
import tempfile
import uuid
from datetime import datetime
from io import BytesIO
from typing import Optional
from uuid import UUID

from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session
from sqlalchemy.orm.attributes import flag_modified

from app.core.storage import StorageBackend, get_storage
from app.domains.tools.modules.pdf_editor.models import (
    CompanyDocument,
    PdfDocument,
)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

ALLOWED_UPLOAD_TYPES = {"application/pdf"}

ALLOWED_CONVERT_TYPES = {
    "application/pdf",
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/webp",
    "image/heic",
    "image/tiff",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


def _hex_to_rgb(hex_color: Optional[str], default=(0.0, 0.0, 0.0)):
    """'#rrggbb' -> (r, g, b) floats in [0, 1]. Falls back to `default` for
    anything else (None, 'transparent', short/invalid hex, ...)."""
    if not hex_color:
        return default
    h = hex_color.lstrip("#")
    if len(h) != 6:
        return default
    try:
        return (
            int(h[0:2], 16) / 255,
            int(h[2:4], 16) / 255,
            int(h[4:6], 16) / 255,
        )
    except ValueError:
        return default


def flatten_annotations_bytes(doc: PdfDocument, storage: StorageBackend) -> bytes:
    """Burn a document's annotations into its PDF pages.

    Pure (doc, storage) -> bytes so callers that already hold the PdfDocument
    ORM object -- e.g. sign_service, which reads the same "linked" document to
    burn signatures -- can flatten without a redundant company-scoped re-fetch.
    Returns the flattened PDF as bytes, or the original bytes unchanged if
    there are no annotations.

    Supports: text, image, stamp, shape (rect/circle/line). Each annotation
    burns independently (one malformed annotation -- e.g. a corrupt image
    data URL -- is skipped, not fatal to the rest of the page).

    NOT supported: freehand "drawing" annotations (pencil tool) and
    "sign_field" placeholders. Fabric.js's exported path data doesn't
    unambiguously say where the path's on-page origin is relative to its own
    point coordinates without either reading fabric's exact version-specific
    internals or live-testing against it -- guessing would risk silently
    misplaced or distorted ink on a document someone is about to sign, which
    is worse than the current "not burned, still visible in the editor"
    behavior. Both types stay visible in the app's own editor/canvas views
    but are silently absent from flattened output, including documents sent
    for e-signature.
    """
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas as rl_canvas

    pdf_data = storage.read(doc.file_path)

    if not doc.annotations:
        return pdf_data

    with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
        local_path = os.path.join(tmpdir, "document.pdf")
        with open(local_path, "wb") as fh:
            fh.write(pdf_data)

        reader = PdfReader(local_path)
        writer = PdfWriter()

        by_page: dict[int, list] = {}
        for ann in doc.annotations:
            pg = ann.get("page", 0)
            by_page.setdefault(pg, []).append(ann)

        for i, page in enumerate(reader.pages):
            page_anns = by_page.get(i, [])
            if page_anns:
                media_box = page.mediabox
                pw = float(media_box.width)
                ph = float(media_box.height)

                packet = BytesIO()
                c = rl_canvas.Canvas(packet, pagesize=(pw, ph))

                drawn = 0
                for ann in page_anns:
                    try:
                        if _burn_one_annotation(c, ann, ph):
                            drawn += 1
                    except Exception:
                        continue  # skip this annotation, keep the rest of the page

                # A Canvas with zero draw calls serializes to a zero-page PDF
                # (reportlab only finalizes a page once something is drawn on
                # it), which would make merge_page below blow up on an empty
                # pages list -- e.g. every annotation on this page was an
                # unsupported/malformed one. Nothing to merge is a no-op, not
                # an error.
                if drawn > 0:
                    c.save()
                    packet.seek(0)
                    overlay_reader = PdfReader(packet)
                    page.merge_page(overlay_reader.pages[0])

            writer.add_page(page)

        flat_path = os.path.join(tmpdir, "flattened.pdf")
        with open(flat_path, "wb") as fh:
            writer.write(fh)
        with open(flat_path, "rb") as fh:
            return fh.read()


def _burn_one_annotation(c, ann: dict, ph: float) -> bool:
    """Draw a single annotation onto the in-progress reportlab canvas `c`.
    Returns whether it actually drew anything -- a Canvas that never draws
    serializes to a zero-page PDF (see the zero-page guard around this
    function's caller), so "didn't raise" alone isn't a safe proxy: several
    branches below intentionally no-op on unusable data (empty image src,
    missing dimensions, unsupported annotation type) without raising.

    `ph` is the page height in points -- every coordinate conversion below
    mirrors it because annotation x/y are stored top-down (matching the
    browser canvas/fabric.js convention) while reportlab draws bottom-up.

    Rotation is intentionally not applied (matches the pre-existing "text"
    burn this was extracted from): fabric.js rotates around an object's
    center by default, and getting the pivot math wrong would visibly
    mis-place the annotation, which is worse than drawing it unrotated.
    """
    ann_type = ann.get("type")
    style = ann.get("style") or {}
    x = ann.get("x", 0)
    y_top = ann.get("y", 0)
    w = ann.get("width") or 0
    h = ann.get("height") or 0

    if ann_type == "text":
        font_size = style.get("font_size") or 12
        r, g, b = _hex_to_rgb(style.get("color"), default=(0.0, 0.0, 0.0))
        c.setFillColorRGB(r, g, b)
        c.setFont("Helvetica", font_size)
        c.drawString(x, ph - y_top - font_size, ann.get("content") or "")
        return True

    if ann_type == "image":
        from reportlab.lib.utils import ImageReader

        content = ann.get("content") or ""
        if not content.startswith("data:") or "," not in content or not w or not h:
            return False
        _, b64data = content.split(",", 1)
        img_reader = ImageReader(BytesIO(base64.b64decode(b64data)))
        opacity = style.get("opacity")
        if opacity is not None and opacity < 1:
            c.saveState()
            c.setFillAlpha(opacity)
            c.setStrokeAlpha(opacity)
        c.drawImage(img_reader, x, ph - y_top - h, w, h, mask="auto")
        if opacity is not None and opacity < 1:
            c.restoreState()
        return True

    if ann_type == "stamp":
        r, g, b = _hex_to_rgb(style.get("color"), default=(0.42, 0.45, 0.50))
        c.setFillColorRGB(r, g, b)
        c.setFont("Helvetica-Bold", 22)
        c.drawString(x, ph - y_top - 22, ann.get("content") or "")
        return True

    if ann_type == "shape":
        stroke_width = style.get("border_width") or 2
        sr, sg, sb = _hex_to_rgb(style.get("color"), default=(0.07, 0.09, 0.16))
        c.setStrokeColorRGB(sr, sg, sb)
        c.setLineWidth(stroke_width)
        fill_hex = style.get("background_color")
        has_fill = bool(fill_hex) and fill_hex != "transparent"
        if has_fill:
            fr, fg, fb = _hex_to_rgb(fill_hex)
            c.setFillColorRGB(fr, fg, fb)

        kind = ann.get("content")
        if kind == "rect" and w and h:
            c.rect(x, ph - y_top - h, w, h, fill=1 if has_fill else 0, stroke=1)
            return True
        if kind == "circle" and w and h:
            c.ellipse(x, ph - y_top - h, x + w, ph - y_top, fill=1 if has_fill else 0, stroke=1)
            return True
        if kind == "line" and w:
            c.line(x, ph - y_top, x + w, ph - y_top)
            return True
        return False

    return False  # unsupported type (drawing, sign_field, ...)


# ===========================================================================
# PdfEditorService
# ===========================================================================

class PdfEditorService:
    """
    Core service for user-owned PDF documents.

    File paths stored in the database are *storage keys* (relative paths),
    e.g. ``{company_id}/pdf-editor/{doc_id}/document.pdf``.
    The active StorageBackend resolves these to local paths or R2 objects.
    """

    def __init__(self, db: Session):
        self.db = db
        self.storage: StorageBackend = get_storage()

    # ------------------------------------------------------------------
    # Key helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _key(company_id: UUID, doc_id: UUID, filename: str) -> str:
        return f"{company_id}/pdf-editor/{doc_id}/{filename}"

    # ------------------------------------------------------------------
    # Local processing helpers
    # ------------------------------------------------------------------

    def _generate_thumbnail(
        self, pdf_local_path: str, thumb_local_path: str
    ) -> bool:
        try:
            from pdf2image import convert_from_path

            images = convert_from_path(
                pdf_local_path, first_page=1, last_page=1, size=(300, None)
            )
            if images:
                images[0].save(thumb_local_path, "PNG")
                return True
        except Exception:
            pass
        return False

    @staticmethod
    def _page_count(pdf_local_path: str) -> int:
        try:
            from pypdf import PdfReader
            return len(PdfReader(pdf_local_path).pages)
        except Exception:
            return 1

    # ------------------------------------------------------------------
    # File-format converters
    # ------------------------------------------------------------------

    def _image_to_pdf(
        self, image_path: str, output_path: str, rotation: int = 0
    ) -> None:
        from PIL import Image, ImageOps
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas as rl_canvas

        img = Image.open(image_path)
        img = ImageOps.exif_transpose(img)
        if rotation:
            img = img.rotate(-rotation, expand=True)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")

        temp_jpg = image_path + ".tmp.jpg"
        img.save(temp_jpg, "JPEG", quality=95)
        img_w, img_h = img.size
        img.close()

        page_w, page_h = letter
        scale = min(page_w / img_w, page_h / img_h, 1.0)
        draw_w, draw_h = img_w * scale, img_h * scale
        x = (page_w - draw_w) / 2
        y = (page_h - draw_h) / 2

        c = rl_canvas.Canvas(output_path, pagesize=letter)
        c.drawImage(temp_jpg, x, y, draw_w, draw_h)
        c.save()

        try:
            os.remove(temp_jpg)
        except OSError:
            pass

    def _docx_to_pdf(self, docx_path: str, output_path: str) -> None:
        try:
            from docx import Document as DocxDocument
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import (
                Paragraph,
                SimpleDocTemplate,
                Spacer,
            )

            doc = DocxDocument(docx_path)
            styles = getSampleStyleSheet()
            story = []
            for para in doc.paragraphs:
                if para.text.strip():
                    story.append(Paragraph(para.text, styles["Normal"]))
                    story.append(Spacer(1, 6))

            if not story:
                story.append(
                    Paragraph("(Empty document)", styles["Normal"])
                )

            pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
            pdf_doc.build(story)
        except Exception as exc:
            raise HTTPException(
                status_code=400, detail=f"Failed to convert DOCX: {exc}"
            )

    def _convert_to_pdf(
        self,
        source_path: str,
        output_path: str,
        mime_type: str,
        rotation: int = 0,
    ) -> None:
        if mime_type.startswith("image/"):
            self._image_to_pdf(source_path, output_path, rotation=rotation)
        elif "wordprocessingml" in mime_type:
            self._docx_to_pdf(source_path, output_path)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Cannot convert {mime_type} to PDF",
            )

    # ------------------------------------------------------------------
    # Storage persistence helpers
    # ------------------------------------------------------------------

    def _persist_pdf_and_thumb(
        self,
        company_id: UUID,
        doc_id: UUID,
        pdf_local_path: str,
        tmpdir: str,
    ) -> tuple[str, int, int, Optional[str]]:
        """Upload PDF + thumbnail to storage.

        Returns (pdf_key, file_size, page_count, thumb_key_or_None).
        """
        pdf_key = self._key(company_id, doc_id, "document.pdf")
        with open(pdf_local_path, "rb") as f:
            pdf_data = f.read()
        self.storage.write(pdf_key, pdf_data, "application/pdf")

        file_size = len(pdf_data)
        page_count = self._page_count(pdf_local_path)

        thumb_local = os.path.join(tmpdir, "thumbnail.png")
        thumb_key: Optional[str] = None
        if self._generate_thumbnail(pdf_local_path, thumb_local):
            thumb_key = self._key(company_id, doc_id, "thumbnail.png")
            with open(thumb_local, "rb") as f:
                self.storage.write(thumb_key, f.read(), "image/png")

        return pdf_key, file_size, page_count, thumb_key

    # ------------------------------------------------------------------
    # Document CRUD
    # ------------------------------------------------------------------

    def upload_document(
        self,
        company_id: UUID,
        user_id: UUID,
        file: UploadFile,
        name: Optional[str] = None,
        rotation: int = 0,
    ) -> PdfDocument:
        if file.content_type not in ALLOWED_CONVERT_TYPES:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file.content_type}",
            )

        content = file.file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail="File too large. Maximum size is 50 MB.",
            )

        doc_id = uuid.uuid4()
        is_pdf = file.content_type == "application/pdf"

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            if is_pdf:
                pdf_path = os.path.join(tmpdir, "document.pdf")
                with open(pdf_path, "wb") as fh:
                    fh.write(content)
            else:
                ext = (
                    file.filename.rsplit(".", 1)[-1]
                    if file.filename and "." in file.filename
                    else "bin"
                )
                original_path = os.path.join(tmpdir, f"original.{ext}")
                with open(original_path, "wb") as fh:
                    fh.write(content)
                pdf_path = os.path.join(tmpdir, "document.pdf")
                self._convert_to_pdf(
                    original_path, pdf_path, file.content_type, rotation
                )

            pdf_key, file_size, page_count, thumb_key = (
                self._persist_pdf_and_thumb(
                    company_id, doc_id, pdf_path, tmpdir
                )
            )

        doc = PdfDocument(
            id=doc_id,
            company_id=company_id,
            created_by=user_id,
            name=name or file.filename or "Untitled",
            file_path=pdf_key,
            file_size=file_size,
            page_count=page_count,
            mime_type="application/pdf",
            source_type="convert" if not is_pdf else "upload",
            annotations=[],
            thumbnail_path=thumb_key,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def get_document(
        self, company_id: UUID, document_id: UUID
    ) -> Optional[PdfDocument]:
        return (
            self.db.query(PdfDocument)
            .filter(
                PdfDocument.id == document_id,
                PdfDocument.company_id == company_id,
                PdfDocument.is_active == True,
            )
            .first()
        )

    def get_document_or_404(
        self, company_id: UUID, document_id: UUID
    ) -> PdfDocument:
        doc = self.get_document(company_id, document_id)
        if not doc:
            raise HTTPException(
                status_code=404, detail="Document not found"
            )
        return doc

    def list_documents(
        self,
        company_id: UUID,
        skip: int = 0,
        limit: int = 20,
        search: Optional[str] = None,
    ) -> tuple[list[PdfDocument], int]:
        query = self.db.query(PdfDocument).filter(
            PdfDocument.company_id == company_id,
            PdfDocument.is_active == True,
        )
        if search:
            query = query.filter(PdfDocument.name.ilike(f"%{search}%"))

        total = query.count()
        items = (
            query.order_by(PdfDocument.created_at.desc())
            .offset(skip)
            .limit(limit)
            .all()
        )
        return items, total

    def rename_document(
        self, company_id: UUID, document_id: UUID, new_name: str
    ) -> PdfDocument:
        doc = self.get_document_or_404(company_id, document_id)
        doc.name = new_name
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def delete_document(
        self, company_id: UUID, document_id: UUID
    ) -> bool:
        doc = self.get_document(company_id, document_id)
        if not doc:
            return False
        doc.is_active = False
        self.db.commit()
        return True

    def duplicate_document(
        self,
        company_id: UUID,
        document_id: UUID,
        user_id: UUID,
        new_name: Optional[str] = None,
    ) -> PdfDocument:
        original = self.get_document_or_404(company_id, document_id)
        new_id = uuid.uuid4()

        # Copy PDF
        new_pdf_key = self._key(company_id, new_id, "document.pdf")
        pdf_data = self.storage.read(original.file_path)
        self.storage.write(new_pdf_key, pdf_data, "application/pdf")

        # Copy thumbnail
        new_thumb_key: Optional[str] = None
        if original.thumbnail_path and self.storage.exists(
            original.thumbnail_path
        ):
            new_thumb_key = self._key(company_id, new_id, "thumbnail.png")
            thumb_data = self.storage.read(original.thumbnail_path)
            self.storage.write(new_thumb_key, thumb_data, "image/png")

        doc = PdfDocument(
            id=new_id,
            company_id=company_id,
            created_by=user_id,
            name=new_name or f"{original.name} (Copy)",
            file_path=new_pdf_key,
            file_size=original.file_size,
            page_count=original.page_count,
            mime_type=original.mime_type,
            source_type="upload",
            annotations=(
                list(original.annotations) if original.annotations else []
            ),
            thumbnail_path=new_thumb_key,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        return doc

    # ------------------------------------------------------------------
    # Page operations
    # ------------------------------------------------------------------

    def merge_documents(
        self,
        company_id: UUID,
        user_id: UUID,
        document_ids: list[str],
        name: Optional[str] = None,
    ) -> PdfDocument:
        from pypdf import PdfReader, PdfWriter

        if len(document_ids) < 2:
            raise HTTPException(
                status_code=400,
                detail="At least 2 documents are required for merge",
            )

        new_id = uuid.uuid4()

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            writer = PdfWriter()
            doc_names: list[str] = []

            for i, did in enumerate(document_ids):
                doc = self.get_document(company_id, UUID(did))
                if not doc:
                    raise HTTPException(
                        status_code=404,
                        detail=f"Document {did} not found",
                    )
                # Download each source PDF to temp
                src_path = os.path.join(tmpdir, f"src_{i}.pdf")
                with open(src_path, "wb") as fh:
                    fh.write(self.storage.read(doc.file_path))
                reader = PdfReader(src_path)
                for page in reader.pages:
                    writer.add_page(page)
                doc_names.append(doc.name)

            merged_path = os.path.join(tmpdir, "document.pdf")
            with open(merged_path, "wb") as fh:
                writer.write(fh)

            pdf_key, file_size, page_count, thumb_key = (
                self._persist_pdf_and_thumb(
                    company_id, new_id, merged_path, tmpdir
                )
            )

        label = name or f"Merged - {', '.join(doc_names[:3])}"
        doc = PdfDocument(
            id=new_id,
            company_id=company_id,
            created_by=user_id,
            name=label,
            file_path=pdf_key,
            file_size=file_size,
            page_count=page_count,
            mime_type="application/pdf",
            source_type="merge",
            annotations=[],
            thumbnail_path=thumb_key,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def reorder_pages(
        self, company_id: UUID, document_id: UUID, page_order: list[int]
    ) -> PdfDocument:
        from pypdf import PdfReader, PdfWriter

        doc = self.get_document_or_404(company_id, document_id)

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            local_path = os.path.join(tmpdir, "document.pdf")
            with open(local_path, "wb") as fh:
                fh.write(self.storage.read(doc.file_path))

            reader = PdfReader(local_path)
            if sorted(page_order) != list(range(len(reader.pages))):
                raise HTTPException(
                    status_code=400,
                    detail="page_order must be a permutation of "
                    "0-indexed page indices",
                )

            writer = PdfWriter()
            for idx in page_order:
                writer.add_page(reader.pages[idx])

            with open(local_path, "wb") as fh:
                writer.write(fh)

            # Persist updated PDF
            with open(local_path, "rb") as fh:
                self.storage.write(
                    doc.file_path, fh.read(), "application/pdf"
                )

            # Regenerate thumbnail
            if doc.thumbnail_path:
                thumb_local = os.path.join(tmpdir, "thumbnail.png")
                if self._generate_thumbnail(local_path, thumb_local):
                    with open(thumb_local, "rb") as fh:
                        self.storage.write(
                            doc.thumbnail_path, fh.read(), "image/png"
                        )

            doc.file_size = os.path.getsize(local_path)

        doc.updated_at = datetime.utcnow()
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def delete_pages(
        self, company_id: UUID, document_id: UUID, page_numbers: list[int]
    ) -> PdfDocument:
        from pypdf import PdfReader, PdfWriter

        doc = self.get_document_or_404(company_id, document_id)

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            local_path = os.path.join(tmpdir, "document.pdf")
            with open(local_path, "wb") as fh:
                fh.write(self.storage.read(doc.file_path))

            reader = PdfReader(local_path)
            total = len(reader.pages)
            to_delete = {p - 1 for p in page_numbers}
            remaining = [i for i in range(total) if i not in to_delete]

            if not remaining:
                raise HTTPException(
                    status_code=400,
                    detail="Cannot delete all pages from a document",
                )

            writer = PdfWriter()
            for idx in remaining:
                writer.add_page(reader.pages[idx])

            with open(local_path, "wb") as fh:
                writer.write(fh)

            with open(local_path, "rb") as fh:
                self.storage.write(
                    doc.file_path, fh.read(), "application/pdf"
                )

            if doc.thumbnail_path:
                thumb_local = os.path.join(tmpdir, "thumbnail.png")
                if self._generate_thumbnail(local_path, thumb_local):
                    with open(thumb_local, "rb") as fh:
                        self.storage.write(
                            doc.thumbnail_path, fh.read(), "image/png"
                        )

            doc.page_count = len(remaining)
            doc.file_size = os.path.getsize(local_path)

        doc.updated_at = datetime.utcnow()
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def rotate_pages(
        self, company_id: UUID, document_id: UUID, rotations: dict[str, int]
    ) -> PdfDocument:
        from pypdf import PdfReader, PdfWriter

        doc = self.get_document_or_404(company_id, document_id)

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            local_path = os.path.join(tmpdir, "document.pdf")
            with open(local_path, "wb") as fh:
                fh.write(self.storage.read(doc.file_path))

            reader = PdfReader(local_path)
            writer = PdfWriter()

            for i, page in enumerate(reader.pages):
                key = str(i + 1)
                if key in rotations:
                    degrees = rotations[key]
                    if degrees not in (90, 180, 270):
                        raise HTTPException(
                            status_code=400,
                            detail=f"Invalid rotation {degrees} for "
                            f"page {key}. Use 90, 180, or 270.",
                        )
                    page.rotate(degrees)
                writer.add_page(page)

            with open(local_path, "wb") as fh:
                writer.write(fh)

            with open(local_path, "rb") as fh:
                self.storage.write(
                    doc.file_path, fh.read(), "application/pdf"
                )

            if doc.thumbnail_path:
                thumb_local = os.path.join(tmpdir, "thumbnail.png")
                if self._generate_thumbnail(local_path, thumb_local):
                    with open(thumb_local, "rb") as fh:
                        self.storage.write(
                            doc.thumbnail_path, fh.read(), "image/png"
                        )

            doc.file_size = os.path.getsize(local_path)

        doc.updated_at = datetime.utcnow()
        self.db.commit()
        self.db.refresh(doc)
        return doc

    # ------------------------------------------------------------------
    # Annotations
    # ------------------------------------------------------------------

    def save_annotations(
        self, company_id: UUID, document_id: UUID, annotations: list
    ) -> PdfDocument:
        doc = self.get_document_or_404(company_id, document_id)
        doc.annotations = annotations
        flag_modified(doc, "annotations")
        doc.updated_at = datetime.utcnow()
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def save_field_definitions(
        self, company_id: UUID, document_id: UUID, fields: list
    ) -> PdfDocument:
        """Persist the reusable field-mapping template for a document.
        Every future SignRequest created against this document copies
        these definitions as its starting sign_fields snapshot."""
        from app.domains.tools.modules.pdf_editor.field_registry import (
            is_valid_source,
        )

        keys = [f["key"] for f in fields]
        if len(keys) != len(set(keys)):
            raise HTTPException(
                status_code=400,
                detail="Field keys must be unique within a document",
            )
        for f in fields:
            binding = f.get("data_binding") or {}
            if binding.get("mode") == "prefilled" and not is_valid_source(
                binding.get("source_entity"), binding.get("source_field")
            ):
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Field '{f.get('label') or f['key']}' is set to "
                        "auto-fill but has no valid data source"
                    ),
                )

        doc = self.get_document_or_404(company_id, document_id)
        doc.field_definitions = fields
        flag_modified(doc, "field_definitions")
        doc.updated_at = datetime.utcnow()
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def flatten_annotations(
        self, company_id: UUID, document_id: UUID
    ) -> bytes:
        """Burn text annotations into the PDF pages.

        Returns the flattened PDF as bytes. If no annotations exist the
        original PDF bytes are returned unchanged.
        """
        doc = self.get_document_or_404(company_id, document_id)
        return flatten_annotations_bytes(doc, self.storage)

    # ------------------------------------------------------------------
    # Multi-image -> PDF
    # ------------------------------------------------------------------

    def images_to_pdf(
        self,
        company_id: UUID,
        user_id: UUID,
        files: list[UploadFile],
        name: Optional[str] = None,
    ) -> PdfDocument:
        from PIL import Image, ImageOps
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas as rl_canvas

        doc_id = uuid.uuid4()

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            pdf_path = os.path.join(tmpdir, "document.pdf")
            page_w, page_h = letter
            c = rl_canvas.Canvas(pdf_path, pagesize=letter)
            page_count = 0

            for upload_file in files:
                content = upload_file.file.read()
                if not content:
                    continue

                img = Image.open(BytesIO(content))
                img = ImageOps.exif_transpose(img)
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")

                temp_img = os.path.join(tmpdir, f"_tmp_{page_count}.jpg")
                img.save(temp_img, "JPEG", quality=95)
                img_w, img_h = img.size
                img.close()

                scale = min(page_w / img_w, page_h / img_h, 1.0)
                draw_w, draw_h = img_w * scale, img_h * scale
                x = (page_w - draw_w) / 2
                y = (page_h - draw_h) / 2

                if page_count > 0:
                    c.showPage()
                c.drawImage(temp_img, x, y, draw_w, draw_h)
                page_count += 1

            if page_count == 0:
                raise HTTPException(
                    status_code=400, detail="No valid images provided"
                )

            c.save()

            pdf_key, file_size, pc, thumb_key = (
                self._persist_pdf_and_thumb(
                    company_id, doc_id, pdf_path, tmpdir
                )
            )

        doc = PdfDocument(
            id=doc_id,
            company_id=company_id,
            created_by=user_id,
            name=name or "Photos PDF",
            file_path=pdf_key,
            file_size=file_size,
            page_count=page_count,
            mime_type="application/pdf",
            source_type="convert",
            annotations=[],
            thumbnail_path=thumb_key,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        return doc

    # ------------------------------------------------------------------
    # Import helpers (estimate / invoice / company-doc)
    # ------------------------------------------------------------------

    def import_estimate_as_document(
        self,
        company_id: UUID,
        user_id: UUID,
        estimate_id: str,
        template: str = "classic",
    ) -> PdfDocument:
        from app.core.pdf_generator import generate_estimate_pdf
        from app.domains.company.models import Company
        from app.domains.estimate.api import _prepare_estimate_pdf_data
        from app.domains.estimate.models import Estimate

        estimate = (
            self.db.query(Estimate)
            .filter(
                Estimate.id == estimate_id,
                Estimate.company_id == company_id,
            )
            .first()
        )
        if not estimate:
            raise HTTPException(
                status_code=404, detail="Estimate not found"
            )

        company = (
            self.db.query(Company)
            .filter(Company.id == company_id)
            .first()
        )
        if not company:
            raise HTTPException(
                status_code=404, detail="Company not found"
            )

        pdf_data = _prepare_estimate_pdf_data(estimate, company, self.db)
        pdf_bytes: bytes = generate_estimate_pdf(pdf_data, template)

        doc_id = uuid.uuid4()

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            pdf_path = os.path.join(tmpdir, "document.pdf")
            with open(pdf_path, "wb") as fh:
                fh.write(pdf_bytes)

            pdf_key, file_size, page_count, thumb_key = (
                self._persist_pdf_and_thumb(
                    company_id, doc_id, pdf_path, tmpdir
                )
            )

        est_number = (
            getattr(estimate, "estimate_number", None) or estimate_id
        )
        doc = PdfDocument(
            id=doc_id,
            company_id=company_id,
            created_by=user_id,
            name=f"Estimate {est_number}",
            file_path=pdf_key,
            file_size=file_size,
            page_count=page_count,
            mime_type="application/pdf",
            source_type="estimate",
            source_id=UUID(estimate_id),
            annotations=[],
            thumbnail_path=thumb_key,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def import_invoice_as_document(
        self,
        company_id: UUID,
        user_id: UUID,
        invoice_id: str,
        template: str = "classic",
    ) -> PdfDocument:
        from app.core.pdf_generator import generate_invoice_pdf
        from app.domains.company.models import Company
        from app.domains.invoice.api import _prepare_invoice_pdf_data
        from app.domains.invoice.models import Invoice

        invoice = (
            self.db.query(Invoice)
            .filter(
                Invoice.id == invoice_id,
                Invoice.company_id == company_id,
            )
            .first()
        )
        if not invoice:
            raise HTTPException(
                status_code=404, detail="Invoice not found"
            )

        company = (
            self.db.query(Company)
            .filter(Company.id == company_id)
            .first()
        )
        if not company:
            raise HTTPException(
                status_code=404, detail="Company not found"
            )

        pdf_data = _prepare_invoice_pdf_data(invoice, company, self.db)
        pdf_bytes: bytes = generate_invoice_pdf(pdf_data, template)

        doc_id = uuid.uuid4()

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            pdf_path = os.path.join(tmpdir, "document.pdf")
            with open(pdf_path, "wb") as fh:
                fh.write(pdf_bytes)

            pdf_key, file_size, page_count, thumb_key = (
                self._persist_pdf_and_thumb(
                    company_id, doc_id, pdf_path, tmpdir
                )
            )

        inv_number = (
            getattr(invoice, "invoice_number", None) or invoice_id
        )
        doc = PdfDocument(
            id=doc_id,
            company_id=company_id,
            created_by=user_id,
            name=f"Invoice {inv_number}",
            file_path=pdf_key,
            file_size=file_size,
            page_count=page_count,
            mime_type="application/pdf",
            source_type="invoice",
            source_id=UUID(invoice_id),
            annotations=[],
            thumbnail_path=thumb_key,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def get_or_create_linked_document(
        self,
        company_id: UUID,
        user_id: UUID,
        company_document_id: str,
    ) -> PdfDocument:
        """Get the persistent, e-signature-capable PdfDocument copy linked
        to a CompanyDocument, creating it on first use via
        import_company_document(). Unlike that method's always-fresh-copy
        semantics (used by the editor's one-off "Import" action), this
        reuses the same linked copy on every call -- which is what makes
        field mapping done once from the company document library stay
        reusable rather than being redone on every visit."""
        existing = (
            self.db.query(PdfDocument)
            .filter(
                PdfDocument.company_id == company_id,
                PdfDocument.source_type == "company_doc",
                PdfDocument.source_id == UUID(company_document_id),
                PdfDocument.is_active == True,
            )
            .order_by(PdfDocument.created_at.desc())
            .first()
        )
        if existing:
            return existing
        return self.import_company_document(
            company_id, user_id, company_document_id
        )

    def import_company_document(
        self,
        company_id: UUID,
        user_id: UUID,
        company_document_id: str,
    ) -> PdfDocument:
        company_doc = (
            self.db.query(CompanyDocument)
            .filter(
                CompanyDocument.id == company_document_id,
                CompanyDocument.company_id == company_id,
                CompanyDocument.is_active == True,
            )
            .first()
        )
        if not company_doc:
            raise HTTPException(
                status_code=404, detail="Company document not found"
            )

        new_id = uuid.uuid4()

        # Copy PDF
        new_pdf_key = self._key(company_id, new_id, "document.pdf")
        pdf_data = self.storage.read(company_doc.file_path)
        self.storage.write(new_pdf_key, pdf_data, "application/pdf")

        # Copy thumbnail
        new_thumb_key: Optional[str] = None
        if company_doc.thumbnail_path and self.storage.exists(
            company_doc.thumbnail_path
        ):
            new_thumb_key = self._key(company_id, new_id, "thumbnail.png")
            thumb_data = self.storage.read(company_doc.thumbnail_path)
            self.storage.write(new_thumb_key, thumb_data, "image/png")

        doc = PdfDocument(
            id=new_id,
            company_id=company_id,
            created_by=user_id,
            name=company_doc.name,
            file_path=new_pdf_key,
            file_size=company_doc.file_size,
            page_count=company_doc.page_count,
            mime_type=company_doc.mime_type,
            source_type="company_doc",
            source_id=UUID(company_document_id),
            annotations=[],
            thumbnail_path=new_thumb_key,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        # Bump use count on source
        company_doc.use_count = (company_doc.use_count or 0) + 1
        company_doc.last_used_at = datetime.utcnow()
        self.db.commit()

        return doc


# ===========================================================================
# CompanyDocumentService
# ===========================================================================

class CompanyDocumentService:
    """
    Service for the company-wide document library.

    Documents here are shared templates / reference files that any
    member of a company can import into their own working documents
    via PdfEditorService.import_company_document().
    """

    def __init__(self, db: Session):
        self.db = db
        self.storage: StorageBackend = get_storage()

    # ------------------------------------------------------------------
    # Key helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _key(company_id: UUID, doc_id: UUID, filename: str) -> str:
        return f"{company_id}/company-documents/{doc_id}/{filename}"

    # ------------------------------------------------------------------
    # Local processing helpers
    # ------------------------------------------------------------------

    def _generate_thumbnail(
        self, pdf_local_path: str, thumb_local_path: str
    ) -> bool:
        try:
            from pdf2image import convert_from_path

            images = convert_from_path(
                pdf_local_path,
                first_page=1,
                last_page=1,
                size=(300, None),
            )
            if images:
                images[0].save(thumb_local_path, "PNG")
                return True
        except Exception:
            pass
        return False

    @staticmethod
    def _page_count(pdf_local_path: str) -> int:
        try:
            from pypdf import PdfReader
            return len(PdfReader(pdf_local_path).pages)
        except Exception:
            return 1

    # ------------------------------------------------------------------
    # CRUD
    # ------------------------------------------------------------------

    def upload(
        self,
        company_id: UUID,
        user_id: UUID,
        file: UploadFile,
        name: str,
        description: Optional[str] = None,
        category: Optional[str] = None,
        tags: Optional[list[str]] = None,
    ) -> CompanyDocument:
        if file.content_type not in ALLOWED_UPLOAD_TYPES:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Only PDF files are accepted. "
                    f"Got: {file.content_type}"
                ),
            )

        content = file.file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail="File too large. Maximum size is 50 MB.",
            )

        doc_id = uuid.uuid4()

        with tempfile.TemporaryDirectory(prefix="scopit_") as tmpdir:
            pdf_path = os.path.join(tmpdir, "document.pdf")
            with open(pdf_path, "wb") as fh:
                fh.write(content)

            page_count = self._page_count(pdf_path)
            file_size = len(content)

            # Persist PDF
            pdf_key = self._key(company_id, doc_id, "document.pdf")
            self.storage.write(pdf_key, content, "application/pdf")

            # Thumbnail
            thumb_key: Optional[str] = None
            thumb_local = os.path.join(tmpdir, "thumbnail.png")
            if self._generate_thumbnail(pdf_path, thumb_local):
                thumb_key = self._key(
                    company_id, doc_id, "thumbnail.png"
                )
                with open(thumb_local, "rb") as fh:
                    self.storage.write(
                        thumb_key, fh.read(), "image/png"
                    )

        doc = CompanyDocument(
            id=doc_id,
            company_id=company_id,
            uploaded_by=user_id,
            name=name,
            description=description,
            file_path=pdf_key,
            file_size=file_size,
            mime_type="application/pdf",
            page_count=page_count,
            category=category,
            tags=tags or [],
            use_count=0,
            thumbnail_path=thumb_key,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def get(
        self, company_id: UUID, doc_id: UUID
    ) -> Optional[CompanyDocument]:
        return (
            self.db.query(CompanyDocument)
            .filter(
                CompanyDocument.id == doc_id,
                CompanyDocument.company_id == company_id,
                CompanyDocument.is_active == True,
            )
            .first()
        )

    def get_or_404(
        self, company_id: UUID, doc_id: UUID
    ) -> CompanyDocument:
        doc = self.get(company_id, doc_id)
        if not doc:
            raise HTTPException(
                status_code=404, detail="Company document not found"
            )
        return doc

    def list(
        self,
        company_id: UUID,
        skip: int = 0,
        limit: int = 20,
        search: Optional[str] = None,
        category: Optional[str] = None,
    ) -> tuple[list[CompanyDocument], int]:
        query = self.db.query(CompanyDocument).filter(
            CompanyDocument.company_id == company_id,
            CompanyDocument.is_active == True,
        )
        if search:
            query = query.filter(
                CompanyDocument.name.ilike(f"%{search}%")
            )
        if category:
            query = query.filter(CompanyDocument.category == category)

        total = query.count()
        items = (
            query.order_by(CompanyDocument.created_at.desc())
            .offset(skip)
            .limit(limit)
            .all()
        )
        return items, total

    def update(
        self,
        company_id: UUID,
        doc_id: UUID,
        name: Optional[str] = None,
        description: Optional[str] = None,
        category: Optional[str] = None,
        tags: Optional[list[str]] = None,
    ) -> CompanyDocument:
        doc = self.get_or_404(company_id, doc_id)
        if name is not None:
            doc.name = name
        if description is not None:
            doc.description = description
        if category is not None:
            doc.category = category
        if tags is not None:
            doc.tags = tags
            flag_modified(doc, "tags")
        doc.updated_at = datetime.utcnow()
        self.db.commit()
        self.db.refresh(doc)
        return doc

    def delete(self, company_id: UUID, doc_id: UUID) -> bool:
        doc = self.get(company_id, doc_id)
        if not doc:
            return False
        doc.is_active = False
        doc.updated_at = datetime.utcnow()
        self.db.commit()
        return True

    def increment_use_count(
        self, company_id: UUID, doc_id: UUID
    ) -> None:
        doc = self.get(company_id, doc_id)
        if doc:
            doc.use_count = (doc.use_count or 0) + 1
            doc.last_used_at = datetime.utcnow()
            self.db.commit()
